"""
Generate x25108093-CEAI-Innovator-Critique-Briefing.docx from the markdown briefing.
Uses python-docx. Run: python generate_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

# ── helpers ──────────────────────────────────────────────────────────────────

def set_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 0:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    elif level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x00, 0xA8, 0x96)
    return h

def add_para(doc, text, bold=False, italic=False, size=10, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_kv_table(doc, rows):
    """Two-column key-value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.0)
    return table

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_section_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('─' * 70)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Normal style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ── COVER PAGE ───────────────────────────────────────────────────────────────
set_heading(doc, 'Innovator Critique: From Pipeline to Position', 0)
set_heading(doc, 'Briefing Document', 1)

add_kv_table(doc, [
    ('Module',      'Customer Engagement and Artificial Intelligence (H9CEAI)'),
    ('Programme',   'MSCAIBUS1 / MSCAI'),
    ('Student ID',  'x25108093'),
    ('Lecturer',    'Victor del Rosal'),
    ('Date',        '14 April 2026'),
    ('Venture',     'MangoMind — AI-Powered Employee Burnout Risk Intelligence for Irish SMBs'),
])
doc.add_paragraph()

# Pipeline parameters
set_heading(doc, 'Pipeline Parameters', 2)
add_para(doc,
    'Note: The AI Innovator interface requires interactive JavaScript rendering to reveal the pipeline prompt. '
    'I resolved this ambiguity by selecting my own parameters, as the brief permits: '
    '"If any instruction in this brief is unclear, state the doubt explicitly and make your own reasoned decision."',
    italic=True, size=9)
doc.add_paragraph()

param_table = doc.add_table(rows=6, cols=2)
param_table.style = 'Table Grid'
params = [('Country','Ireland'),('Fruit','Mango'),('Animal','Dolphin'),
          ('Action','Track'),('Wild Card (5th)','Burnout'),('Venture Name','MangoMind')]
for i,(k,v) in enumerate(params):
    param_table.rows[i].cells[0].text = k
    param_table.rows[i].cells[1].text = v
    param_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    shade_cell(param_table.rows[i].cells[0], 'F2F2F2')
doc.add_paragraph()

set_heading(doc, 'Repository & URLs', 2)
add_kv_table(doc, [
    ('GitHub Repository',        'https://github.com/Vinayakcsnci/x25108093-innovator-critique'),
    ('Landing Page (Pages)',     'https://vinayakcsnci.github.io/x25108093-innovator-critique/index.html'),
    ('Pitch Deck (Pages)',       'https://vinayakcsnci.github.io/x25108093-innovator-critique/pitch.html'),
    ('Prototype (Pages)',        'https://vinayakcsnci.github.io/x25108093-innovator-critique/prototype.html'),
    ('Initial Commit Hash',      '5e7713159ea958c32bd6c29b2f99c2d78d5a2c4f'),
    ('Prototype Commit Hash',    '5e7713159ea958c32bd6c29b2f99c2d78d5a2c4f (included in initial push)'),
])
doc.add_page_break()

# ── PART A ────────────────────────────────────────────────────────────────────
set_heading(doc, 'Part A: Regulatory Audit of Own Artefacts', 1)
add_para(doc, 'Artefacts audited: MangoMind/index.html and MangoMind/pitch.html', italic=True, size=9)
add_para(doc, 'Regulatory frameworks: EU AI Act, GDPR, Irish Consumer Protection Act 2007 (CPA), Irish employment law.')
doc.add_paragraph()

# Part A table
a_headers = ['Artefact','Issue','Severity','Quoted Offending Text','Regulatory Basis','Recommended Fix']
a_rows = [
    ['index.html','Unsubstantiated "3 weeks before" detection claim — no clinical trial evidence','High',
     '"Know your team is burning out 3 weeks before they do"',
     'EU AI Act Art. 13; CPA 2007 s.43',
     'Qualify: "based on beta pilot data, n=5 teams. Clinical validation with UCD in progress Q3 2026"'],
    ['pitch.html','€11B employer cost overstates direct business cost by ~3× (actual: €3.7B employer-specific)','High',
     '"€11B — Annual cost of workplace mental health to Irish businesses"',
     'EU AI Act Art. 13; CPA 2007 s.43',
     'Change to €3.7B citing Mental Health Reform "Counting the Cost" (2022)'],
    ['pitch.html','78% IBEC statistic is unverifiable — hallucinated source attribution','High',
     '"78% of burnout-related departures… (IBEC 2023)"',
     'CPA 2007 s.43; GDPR Art. 5(1)(d) accuracy principle',
     'Replace with verifiable figure: "54% of HR managers use reactive-only approach (IBEC Workplace Wellbeing Survey, 2023)"'],
    ['index.html','DPIA asserted without published summary, DPO contact, or DPC consultation reference','High',
     '"fully GDPR-compliant and operates under a completed DPIA"',
     'GDPR Art. 13, 35; Irish DPC Guidance',
     'Publish DPIA summary; name DPO; state DPC consultation status per Art. 36'],
    ['index.html','EU AI Act limited-risk self-classification not legally verified; Annex III (4)(a) may apply','High',
     '"MangoMind is classified as a limited-risk AI system"',
     'EU AI Act Annex III (4)(a); Art. 6',
     'Obtain formal legal opinion; state "classification review in progress with legal counsel"'],
    ['pitch.html','94% retention on n=5 is arithmetically incoherent (4.7 customers)','Medium',
     '"94% Beta customer retention at 90 days"',
     'Irish CPA 2007 s.42; CCPC testimonial guidance',
     'Change to "80% retention (4/5 beta customers, n=5 — early signal only)"'],
    ['index.html','"ISO 27001 certification in progress" may be read as near-complete','Medium',
     '"ISO 27001 certification in progress"',
     'CPA 2007 s.43 (misleading omission)',
     'Rephrase: "pursuing ISO 27001; expected Q4 2026"'],
    ['index.html','No accessibility statement for a health-adjacent product','Low',
     'Landing page — no WCAG notice present',
     'SI 358/2020 (Web Accessibility); Irish Disability Act 2005 s.26',
     'Add WCAG 2.1 AA conformance statement'],
]

tbl_a = doc.add_table(rows=1 + len(a_rows), cols=6)
tbl_a.style = 'Table Grid'
widths = [Inches(0.8), Inches(1.5), Inches(0.65), Inches(1.5), Inches(1.0), Inches(1.3)]
for j, (hdr, w) in enumerate(zip(a_headers, widths)):
    cell = tbl_a.rows[0].cells[j]
    cell.text = hdr
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(8)
    shade_cell(cell, '1A1A2E')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

SEV_COLORS = {'High':'FFD0D0','Medium':'FFF3CC','Low':'D6F0E0'}
for i, row_data in enumerate(a_rows):
    for j, val in enumerate(row_data):
        cell = tbl_a.rows[i+1].cells[j]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        if j == 2:
            shade_cell(cell, SEV_COLORS.get(val, 'FFFFFF'))
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()
set_heading(doc, 'Regulatory Risk Posture', 2)
add_para(doc,
    'MangoMind\'s artefacts carry material legal exposure across three vectors. The EU AI Act self-classification as '
    'limited-risk is the most urgent: if Annex III (4)(a) applies to workplace wellbeing monitoring in an employment '
    'context — which legal counsel must determine before launch — deployment without a conformity assessment constitutes '
    'non-compliance as of August 2024. Separately, three statistics in the artefacts are unverifiable or arithmetically '
    'incoherent, creating Irish consumer law exposure under the CPA 2007. The GDPR DPIA assertion without any supporting '
    'disclosure fails Article 13 obligations. None of these issues are fatal, but none should survive into a commercial '
    'launch without correction.',
    size=10)

doc.add_page_break()

# ── PART B ────────────────────────────────────────────────────────────────────
set_heading(doc, 'Part B: Trust Transfer Audit', 1)
add_para(doc, 'Auditor persona: CIPD-qualified HR Manager, 65-person Irish tech startup, personally affected by burnout-related team departure in 2024.', italic=True, size=9)
doc.add_paragraph()

b_moments = [
    ('Moment 1 — Precision without evidence',
     '"Know your team is burning out 3 weeks before they do"',
     'MangoMind/index.html (hero headline)',
     'An HR manager with clinical literacy knows burnout onset is a process, not a discrete event. '
     'The claim of 3-week advance detection without a single cited study, trial, or methodology reads not as a '
     'bold value proposition but as hallucinated precision. The specific number ("3 weeks") signals either that '
     'the founders have clinical evidence they have not published, or that the number was generated by a pipeline '
     'and no one pushed back on it. An HR manager who has sat through a real burnout crisis knows the difference — '
     'and the absence of evidence transforms the headline from inspiring to alarming.'),
    ('Moment 2 — Arithmetic failure in traction data',
     '"94% Beta customer retention at 90 days"',
     'MangoMind/pitch.html (Slide 8: Traction)',
     'The same slide states "5 beta customers actively using MangoMind." A 94% retention rate calculated on n=5 '
     'customers at 90 days — where 100% is 5 customers and 94% is 4.7 — is a statistic that does not survive '
     'arithmetic scrutiny. The moment she reaches for her calculator is the moment trust exits the room. A metric '
     'that breaks under basic inspection destroys more confidence than no metric at all.'),
    ('Moment 3 — Compliance assertion without substance',
     '"Our system is fully GDPR-compliant and operates under a completed Data Protection Impact Assessment (DPIA)."',
     'MangoMind/index.html (Privacy & Compliance section)',
     'Before deploying any data monitoring tool to 65 employees, the HR manager\'s legal team will ask: "Can we see '
     'the DPIA summary?" The landing page asserts completion but provides no link, no DPO contact, no DPC consultation '
     'reference. Claiming compliance without any supporting documentation for a tool that monitors worker behaviour is '
     'not reassurance — it is a red flag. The claim manufactures trust and then immediately withdraws the substance '
     'that would support it.'),
]
for title, quote, source, explanation in b_moments:
    set_heading(doc, title, 2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f'"{quote}"')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)
    add_para(doc, f'File: {source}', italic=True, size=9)
    add_para(doc, explanation, size=10)
    doc.add_paragraph()

doc.add_page_break()

# ── PART C ────────────────────────────────────────────────────────────────────
set_heading(doc, 'Part C: Research and Correct', 1)
add_para(doc, 'Files updated: MangoMind/index.html, MangoMind/pitch.html', italic=True, size=9)
doc.add_paragraph()
add_para(doc, 'Research methodology: Independent web search (2025–2026 sources), Mental Health Reform, IBEC, CIPD Ireland, EU AI Act official text (EUR-Lex), and occupational psychology literature.', size=9)
doc.add_paragraph()

c_headers = ['Error','Original Claim','Corrected Claim','Severity']
c_rows = [
    ['€11B employer cost hallucination',
     '"€11B annual cost to Irish businesses"',
     '"€3.7B annual cost to Irish employers (Mental Health Reform, Counting the Cost, 2022)"','High'],
    ['IBEC 78% hallucinated statistic',
     '"78% of departures, warning signs unacted (IBEC 2023)"',
     '"54% of HR managers use reactive-only approach (IBEC Workplace Wellbeing Survey, 2023)"','High'],
    ['"3 weeks before" unqualified claim',
     '"Know your team is burning out 3 weeks before they do"',
     'Added: "based on beta pilot data, n=5 teams. Clinical validation with UCD in progress Q3 2026"','High'],
    ['EU AI Act self-declaration',
     '"classified as a limited-risk AI system"',
     '"designed to qualify as limited-risk; formal classification review in progress with legal counsel"','High'],
    ['€14K CIPD attribution inaccurate',
     '"€14K average cost (CIPD Ireland 2024)"',
     '"€10K–€14K estimated, based on CIPD UK benchmarks adjusted for Irish market, 2024"','Medium'],
    ['94% retention metric incoherent',
     '"94% Beta customer retention at 90 days"',
     '"80% retention (4 of 5 beta customers, n=5 cohort — early signal only)"','Medium'],
]
tbl_c = doc.add_table(rows=1+len(c_rows), cols=4)
tbl_c.style = 'Table Grid'
for j, hdr in enumerate(c_headers):
    cell = tbl_c.rows[0].cells[j]
    cell.text = hdr
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F0F0F0')
for i, row_data in enumerate(c_rows):
    for j, val in enumerate(row_data):
        cell = tbl_c.rows[i+1].cells[j]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        if j == 3:
            shade_cell(cell, SEV_COLORS.get(val, 'FFFFFF'))
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()
set_heading(doc, 'Do These Errors Significantly Alter the Project?', 2)
add_para(doc,
    'The core proposition — anonymised, team-level burnout risk detection for Irish SMBs — is commercially coherent '
    'and addresses a real problem. The statistical hallucinations inflate the market case but do not fabricate it. '
    'The EU AI Act classification issue is the most material: a high-risk determination would add 6–12 months to '
    'the compliance roadmap and reshape capital requirements. The 94% retention metric is the most immediately '
    'damaging — it reveals that the pipeline generated a percentage from an incoherent base (n=5). An investor '
    'reading closely will catch this and lose confidence in all other numbers.',
    size=10)

doc.add_page_break()

# ── PART D ────────────────────────────────────────────────────────────────────
set_heading(doc, 'Part D: Prototype', 1)
add_para(doc, 'File: prototype.html (repository root) — served via GitHub Pages at:', size=9, italic=True)
add_para(doc, 'https://vinayakcsnci.github.io/x25108093-innovator-critique/prototype.html', size=9)
doc.add_paragraph()
add_para(doc,
    'MangoMind\'s alpha prototype demonstrates the core value proposition: an HR manager sees a team-level Burnout '
    'Risk Index (BRI) computed in real time from four signal-category sliders (after-hours messaging, meeting density, '
    'pulse check-in sentiment, response latency), colour-coded from green (healthy, BRI ≤ 30) to red (critical, BRI > 80).',
    size=10)
doc.add_paragraph()
set_heading(doc, 'Core Interactions Demonstrated', 2)
for feat in [
    'Live BRI gauge (SVG) — real-time weighted formula: after-hours 35%, meetings 25%, sentiment 25%, latency 15%',
    'Signal-specific recommendations — the dominant contributing signal drives a tailored HR action message (not just a generic BRI band label)',
    'Signal breakdown table — shows each signal\'s contribution as a percentage of the total BRI',
    'Trend vs. 4-week baseline indicator (improving / stable / worsening)',
    'Signal-specific action buttons relevant to the highest-risk driver',
    '4-week historical sparkline showing simulated trend',
    'Transparency panel: exact list of what MangoMind tracks (✓) and does not track (✗) including "Message content — never read"',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(feat).font.size = Pt(10)
doc.add_paragraph()
add_para(doc, 'Commit hash: 5e7713159ea958c32bd6c29b2f99c2d78d5a2c4f', bold=True, size=9)

doc.add_page_break()

# ── PART E ────────────────────────────────────────────────────────────────────
set_heading(doc, 'Part E: Business Modelling', 1)
set_heading(doc, 'Position: Ship with Named Conditions (Option b)', 2)
doc.add_paragraph()
add_para(doc,
    'MangoMind should not ship in 90 days in its current form, but it is not a "do not ship" situation. The product '
    'has a real problem, a coherent solution, and early market signal. Three non-negotiable conditions must be met '
    'before launch.',
    size=10)
doc.add_paragraph()

set_heading(doc, 'Advisory Team Stress-Test', 2)
advisors = [
    ('Dr. Alison Walsh, employment law barrister',
     'Your EU AI Act classification is self-declared. Systems monitoring worker behaviour in an employment context '
     'to infer health states are squarely in the zone of Annex III (4)(a). You have not conducted a conformity '
     'assessment. Do not ship until this is legally settled.'),
    ('Ciarán Hayes, CFO',
     'Your 94% retention metric is computed on n=5. No institutional investor will accept this as a traction metric. '
     'You\'ve also cited a statistic (IBEC 78%) that cannot be verified. Fix your numbers before you close this round.'),
    ('Prof. Mary Dempsey, occupational psychologist',
     'Your clinical claim — detect burnout 3 weeks ahead — has no RCT behind it. Deploying this at scale without a '
     'validation study risks causing harm: HR managers acting on false positives. The UCD study is the right call. '
     'Do not make clinical efficacy claims until you have the data.'),
    ('Siobhán Ní Mhurchú, trade union official',
     'No Irish trade union has been consulted on this technology. The moment a union receives a complaint that '
     'management is using AI to monitor worker stress without negotiation, your reputation is finished. Engage '
     'unions now, before customers sign.'),
]
for name, quote in advisors:
    set_heading(doc, name, 2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f'"{quote}"')
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    doc.add_paragraph()

set_heading(doc, 'Pre-Mortem: Three Failure Modes', 2)
for fm in [
    'Regulatory misclassification — shipping under incorrect AI Act risk category triggers DPC investigation',
    'Trust betrayal at first adverse event — any customer misuse of BRI data becomes MangoMind\'s reputational liability',
    'Clinical overclaiming — a false-positive BRI alert leads to documented HR action; employee suffers harm; negligence exposure follows',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(fm).font.size = Pt(10)
doc.add_paragraph()

set_heading(doc, 'The Three Non-Negotiable Conditions', 2)
conditions = [
    ('Condition 1: EU AI Act conformity assessment',
     'Without a formal legal opinion confirming limited-risk status, shipping is regulatory non-compliance. Timeline: 90 days. Cost: €25K.'),
    ('Condition 2: IBEC/ICTU joint engagement; model union clause agreed',
     'Irish employment law requires consultation with employee representatives before deploying behavioural monitoring. Timeline: 60–90 days. Cost: €10K.'),
    ('Condition 3: All artefact statistics independently verified and corrected',
     'Three claims are currently unverifiable or wrong (Part C). An investor detecting fabricated statistics during diligence will withdraw. Timeline: 2 weeks. Cost: zero.'),
]
for title, body in conditions:
    set_heading(doc, title, 2)
    add_para(doc, body, size=10)
    doc.add_paragraph()

set_heading(doc, 'Scenario Analysis', 2)
sc_headers = ['Scenario','Year 1 ARR','Year 3 ARR','Key Assumption']
sc_rows = [
    ['Best case','€720K','€9.2M','IBEC endorsement + EAP partnership in Year 1'],
    ['Base case','€450K','€6.3M','25 → 120 → 280 customers as per plan'],
    ['Worst case','€180K','€2.1M','AI Act high-risk classification delays launch 6 months'],
]
tbl_s = doc.add_table(rows=1+len(sc_rows), cols=4)
tbl_s.style = 'Table Grid'
for j, hdr in enumerate(sc_headers):
    cell = tbl_s.rows[0].cells[j]
    cell.text = hdr
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F0F0F0')
for i, row in enumerate(sc_rows):
    for j, val in enumerate(row):
        tbl_s.rows[i+1].cells[j].text = val
        tbl_s.rows[i+1].cells[j].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_page_break()

# ── PART F ────────────────────────────────────────────────────────────────────
set_heading(doc, 'Part F: The Plan — 90-Day and 180-Day Capital Allocation', 1)
add_para(doc, 'Total Seed Capital: €2,000,000 | 90-Day Period: April–June 2026 | 180-Day Period: to September 2026', italic=True, size=9)
doc.add_paragraph()
add_para(doc,
    'The €2M seed must achieve one objective: get MangoMind to a credible Series A conversation. The regulatory and '
    'compliance workstream (Part A findings) is the critical path — it gates every other activity.',
    size=10)
doc.add_paragraph()

set_heading(doc, '90-Day Plan (€274K deployed of €1M first tranche)', 2)

set_heading(doc, 'Hiring Sequence', 2)
hire_rows = [
    ('Day 1','CEO, CTO, CSO (founders)','Founding team already committed'),
    ('Week 2','Legal Counsel (part-time retainer, €15K)','EU AI Act conformity assessment is critical path'),
    ('Week 4','Senior ML Engineer (FTE, €85K pa)','Rebuild BRI model to production-grade'),
    ('Week 6','Head of Product/Design (FTE, €75K pa)','Dashboard UX required for beta-to-paid conversion'),
    ('Week 8','DPO (fractional, €2K/mo)','GDPR DPIA registration and DPC liaison'),
]
tbl_h = doc.add_table(rows=1+len(hire_rows), cols=3)
tbl_h.style = 'Table Grid'
for j, hdr in enumerate(['Week','Role','Rationale']):
    cell = tbl_h.rows[0].cells[j]
    cell.text = hdr
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F0F0F0')
for i, row in enumerate(hire_rows):
    for j, val in enumerate(row):
        tbl_h.rows[i+1].cells[j].text = val
        tbl_h.rows[i+1].cells[j].paragraphs[0].runs[0].font.size = Pt(8)
doc.add_paragraph()

set_heading(doc, '90-Day Line-Item Budget', 2)
budget_90 = [
    ('Salaries (ML Engineer 3mo + Head of Product 1.5mo + founders)','€115K'),
    ('Legal & Compliance (EU AI Act, DPIA, union clause)','€55K'),
    ('Product & Engineering (infrastructure, AWS, tools)','€20K'),
    ('Sales & Marketing (IBEC pipeline, CIPD conference, content)','€48K'),
    ('DPO (fractional, 3 months)','€6K'),
    ('Advisor fees (clinical, legal)','€15K'),
    ('Operations & Admin','€15K'),
    ('Total 90-Day Deployed','€274K'),
    ('Retained for Days 91–180','€726K'),
]
tbl_b90 = doc.add_table(rows=len(budget_90), cols=2)
tbl_b90.style = 'Table Grid'
for i,(cat,amt) in enumerate(budget_90):
    tbl_b90.rows[i].cells[0].text = cat
    tbl_b90.rows[i].cells[1].text = amt
    if i == len(budget_90)-2:
        for cell in tbl_b90.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True
            shade_cell(cell, 'FFF3CC')
    elif i == len(budget_90)-1:
        for cell in tbl_b90.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True
    tbl_b90.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    tbl_b90.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
doc.add_paragraph()

set_heading(doc, '180-Day Plan (additional €543K deployed; €183K held as reserve)', 2)
budget_180 = [
    ('Salaries (3 new hires: AE, CSM, Senior BE + full team)','€295K'),
    ('Sales & Marketing (EAP partnership, PR launch, case studies)','€103K'),
    ('Clinical Validation Study (UCD)','€40K'),
    ('ISO 27001 Audit','€20K'),
    ('Legal (ongoing compliance, partnership contracts)','€25K'),
    ('Infrastructure Scaling','€25K'),
    ('Operations & Admin','€20K'),
    ('Customer Onboarding & Support','€15K'),
    ('Total Days 91–180','€543K'),
    ('Reserve (buffer, strategic optionality)','€183K'),
]
tbl_b180 = doc.add_table(rows=len(budget_180), cols=2)
tbl_b180.style = 'Table Grid'
for i,(cat,amt) in enumerate(budget_180):
    tbl_b180.rows[i].cells[0].text = cat
    tbl_b180.rows[i].cells[1].text = amt
    if i >= len(budget_180)-2:
        for cell in tbl_b180.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True
            shade_cell(cell, 'FFF3CC')
    tbl_b180.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    tbl_b180.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
doc.add_paragraph()

set_heading(doc, 'Full 180-Day Capital Summary', 2)
summary_rows = [
    ('Salaries (8 FTE + fractionals)','€410K','20.5%'),
    ('Legal & Compliance','€80K','4%'),
    ('Sales & Marketing','€151K','7.55%'),
    ('Clinical Research (UCD)','€40K','2%'),
    ('Infrastructure','€45K','2.25%'),
    ('Product & Engineering (tools)','€45K','2.25%'),
    ('Operations, Admin, Advisors','€56K','2.8%'),
    ('Total Deployed','€827K','41.35%'),
    ('Reserve (runway + strategic)','€1,173K','58.65%'),
]
tbl_sum = doc.add_table(rows=1+len(summary_rows), cols=3)
tbl_sum.style = 'Table Grid'
for j, hdr in enumerate(['Category','Amount','% of €2M']):
    cell = tbl_sum.rows[0].cells[j]
    cell.text = hdr
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F0F0F0')
for i, row in enumerate(summary_rows):
    for j, val in enumerate(row):
        tbl_sum.rows[i+1].cells[j].text = val
        tbl_sum.rows[i+1].cells[j].paragraphs[0].runs[0].font.size = Pt(9)
    if row[0] in ('Total Deployed','Reserve (runway + strategic)'):
        for cell in tbl_sum.rows[i+1].cells:
            cell.paragraphs[0].runs[0].bold = True
            shade_cell(cell, 'FFF3CC')

doc.add_paragraph()
set_heading(doc, 'What We Are Deliberately NOT Doing', 2)
not_doing = [
    ('UK expansion before Day 180','Ireland is proof-of-concept. UK entry is a Series A objective.'),
    ('Enterprise (250+ employees) sales','SMB wedge first; enterprise procurement cycles are too long at this stage.'),
    ('Custom AI model from scratch','Fine-tuned BERT + federated learning is sufficient; proprietary training data is Year 2.'),
    ('Acquisitions','No M&A before Series A.'),
    ('PR agency retainer','Direct outreach and freelance PR at €3–5K/campaign vs. €8–12K/month agency.'),
]
tbl_nd = doc.add_table(rows=1+len(not_doing), cols=2)
tbl_nd.style = 'Table Grid'
for j, hdr in enumerate(['Not Doing','Why']):
    cell = tbl_nd.rows[0].cells[j]
    cell.text = hdr
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F0F0F0')
for i,(what,why) in enumerate(not_doing):
    tbl_nd.rows[i+1].cells[0].text = what
    tbl_nd.rows[i+1].cells[1].text = why
    tbl_nd.rows[i+1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    tbl_nd.rows[i+1].cells[1].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_page_break()

# ── REFLECTION ────────────────────────────────────────────────────────────────
set_heading(doc, 'Reflection on AI Use', 1)
add_para(doc,
    'This assessment was completed with the AI Innovator pipeline (for pipeline content generation) and Claude Code '
    '(for file structure, critique drafting, code, and research verification). The pipeline generated a coherent venture '
    'concept rapidly. The critique work — identifying hallucinations, verifying statistics, conducting legal analysis, '
    'and taking a defensible position — was human-led and AI-assisted. The 78% IBEC statistic and the 94% retention '
    'metric were hallucinations that the pipeline produced without flagging; only independent research surfaced them. '
    'This is the trust transfer problem in practice: the artefacts looked professional enough that a non-critical reader '
    'would not question them. Critical thinking, not tooling, determined whether the output was fit for purpose.',
    size=10)

# ── SAVE ─────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), 'x25108093-CEAI-Innovator-Critique-Briefing.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
